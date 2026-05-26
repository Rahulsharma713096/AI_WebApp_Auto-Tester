import json


class JSONReportGenerator:
    def generate(self, data: dict) -> str:
        return json.dumps(data, indent=2, default=str)

    def save(self, data: dict, output_path: str):
        with open(output_path, "w") as f:
            f.write(self.generate(data))


class PDFReportGenerator:
    def generate(self, data: dict) -> bytes:
        from io import BytesIO

        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []

        elements.append(Paragraph(f"Test Report: {data.get('url', 'N/A')}", styles["Title"]))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(f"Score: {data.get('score', 'N/A')}/100", styles["Heading2"]))
        elements.append(Paragraph(f"Duration: {data.get('duration', 'N/A')}s", styles["Normal"]))
        elements.append(Spacer(1, 12))

        if "issues" in data and data["issues"]:
            elements.append(Paragraph("Issues Found:", styles["Heading2"]))
            for issue in data["issues"]:
                elements.append(Paragraph(
                    f"<b>{issue.get('severity', 'info').upper()}:</b> {issue.get('title', '')}",
                    styles["Normal"],
                ))

        doc.build(elements)
        return buffer.getvalue()

    def save(self, data: dict, output_path: str):
        content = self.generate(data)
        with open(output_path, "wb") as f:
            f.write(content)


class DOCXReportGenerator:
    def generate(self, data: dict):
        from docx import Document

        doc = Document()
        doc.add_heading(f"Test Report: {data.get('url', 'N/A')}", 0)

        doc.add_heading("Summary", level=1)
        doc.add_paragraph(f"Score: {data.get('score', 'N/A')}/100")
        doc.add_paragraph(f"Duration: {data.get('duration', 'N/A')}s")
        doc.add_paragraph(f"Status: {data.get('status', 'N/A')}")

        if "issues" in data and data["issues"]:
            doc.add_heading("Issues", level=1)
            for issue in data["issues"]:
                doc.add_paragraph(
                    f"{issue.get('severity', 'info').upper()}: {issue.get('title', '')}"
                )

        return doc

    def save(self, data: dict, output_path: str):
        doc = self.generate(data)
        doc.save(output_path)


def get_generator(format: str = "json"):
    if format == "pdf":
        return PDFReportGenerator()
    elif format == "docx":
        return DOCXReportGenerator()
    return JSONReportGenerator()
